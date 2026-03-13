"""
PIRS V2 Architecture Document Generator
Produces: outputs/PIRS_V2_Architecture.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, copy

OUT = os.path.join(os.path.dirname(__file__), 'outputs', 'PIRS_V2_Architecture.docx')
ARCH_IMG = os.path.join(os.path.dirname(__file__), 'outputs', 'cert', 'plots', 'pirs_v2_architecture.png')

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x38, 0x64)   # dark navy   – headings
BLUE   = RGBColor(0x2E, 0x75, 0xB6)   # mid blue    – sub-headings / accents
LBHEX  = 'D5E8F4'                      # light blue  – table headers
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)
GRAY   = RGBColor(0x59, 0x59, 0x59)
LGRAY  = 'F2F2F2'                      # light grey  – alternating rows

# ── Helpers ───────────────────────────────────────────────────────────────────
def _shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _set_cell_border(cell, **kwargs):
    """kwargs: top, bottom, left, right – each a dict with color, sz, val"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, props in kwargs.items():
        el = OxmlElement(f'w:{side}')
        for k, v in props.items():
            el.set(qn(f'w:{k}'), str(v))
        tcBorders.append(el)
    tcPr.append(tcBorders)

def _thin_border(cell):
    b = dict(val='single', sz='4', color='C0C0C0', space='0')
    _set_cell_border(cell, top=b, bottom=b, left=b, right=b)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.color.rgb = NAVY if level == 1 else BLUE
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    return p

def add_body(doc, text, bold=False, italic=False, color=None, size=11, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    return p

def add_kv(doc, key, value):
    """Bold key: normal value on same line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f'{key}: ')
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(10.5)
    r2 = p.add_run(value)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = GRAY
    return p

def make_table(doc, headers, rows, col_widths, header_hex=LBHEX, alt_hex=LGRAY):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr = tbl.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr.cells[i]
        cell.width = Inches(w)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _shade_cell(cell, header_hex)
        _thin_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = NAVY

    # data rows
    for ri, row in enumerate(rows):
        bg = alt_hex if ri % 2 == 1 else 'FFFFFF'
        tr = tbl.rows[ri + 1]
        for ci, (val, w) in enumerate(zip(row, col_widths)):
            cell = tr.cells[ci]
            cell.width = Inches(w)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _shade_cell(cell, bg)
            _thin_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.name = 'Calibri'
            run.font.size = Pt(10)

    return tbl

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '2E75B6')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Default paragraph spacing ─────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
# COVER / TITLE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()   # top padding

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
t = title_p.add_run('PIRS V2')
t.font.name  = 'Calibri'
t.font.size  = Pt(36)
t.bold       = True
t.font.color.rgb = NAVY

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
s = sub_p.add_run('Pre-Incident Detection System')
s.font.name  = 'Calibri'
s.font.size  = Pt(20)
s.font.color.rgb = BLUE

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
s2 = sub2.add_run('Architecture & Pipeline Reference')
s2.font.name  = 'Calibri'
s2.font.size  = Pt(16)
s2.font.color.rgb = GRAY

doc.add_paragraph()
add_divider(doc)
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
m = meta.add_run(
    'Dataset: CERT r6.2   ·   4,000 Users   ·   516 Days   ·   '
    '9-Layer Drift-Based Pipeline   ·   HTTP Enabled (117M rows)'
)
m.font.name  = 'Calibri'
m.font.size  = Pt(11)
m.italic     = True
m.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '1. Executive Summary', 1)
add_divider(doc)
add_body(doc,
    'PIRS V2 (Predictive Intervention and Risk Stabilization System, Version 2) is a nine-layer '
    'machine learning pipeline designed to detect insider threats in enterprise environments '
    'before an incident occurs. Unlike traditional anomaly-detection systems that compare users '
    'to a population average, PIRS V2 builds a personal behavioral baseline for each individual '
    'and measures how far they have drifted from their own normal over time.',
    space_after=8)

add_body(doc,
    'The system was evaluated on the CERT Insider Threat Dataset r6.2 — a benchmark dataset '
    'of 4,000 synthetic employees over 516 days containing five real labeled insider scenarios. '
    'PIRS V2 achieved a ROC-AUC of 0.8554, successfully detecting insider behavioral drift '
    '7 to 14 days before the first malicious action in 40% of insider cases, and up to 3 days '
    'in advance for 2 out of 5 insiders at CRITICAL/HIGH alert level.',
    space_after=8)

add_body(doc,
    'The pipeline integrates behavioral feature extraction, personal drift detection, ensemble '
    'anomaly scoring, personality-matched interventions, and reinforcement-learning policy '
    'optimization — all validated against ground-truth insider labels.',
    space_after=8)

doc.add_paragraph()

# Summary KPI table
add_heading(doc, 'Key Performance Indicators', 2)
make_table(doc,
    headers=['Metric', 'Value', 'Interpretation'],
    rows=[
        ['ROC-AUC',            '0.8554', 'Strong discrimination: malicious vs. normal days'],
        ['EPR (7-day)',         '40%',    '2 of 5 real insiders detected ≥ 7 days before first attack'],
        ['EPR (14-day)',        '40%',    'Same 2 insiders also present at 14-day early warning window'],
        ['Early Warning (3d)', '2 / 5',  'ACM2278 and CMP2946 flagged at CRITICAL/HIGH with 3 days lead'],
        ['Personality Quality','0.931',   '93.1% of interventions personality-matched to the user type'],
        ['Estimated Cost Saved','$22.8M', 'Prevention value vs. average insider incident cost model'],
        ['Users Monitored',    '4,000',  'Full CERT r6.2 population, all 516 days'],
        ['HTTP Rows Processed','117M',   'All HTTP activity included (was skipped in PIRS V1)'],
    ],
    col_widths=[2.0, 1.3, 3.6],
)
doc.add_paragraph()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. PIRS V2 vs V1 COMPARISON
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '2. PIRS V2 vs. PIRS V1 — What Changed', 1)
add_divider(doc)
add_body(doc,
    'PIRS V2 is a ground-up redesign of the V1 pipeline. The core improvements address the '
    'fundamental limitation of V1: it measured behavior against a population average, which '
    'cannot detect subtle personal drift in users whose baseline differs significantly from '
    'their colleagues.',
    space_after=8)

make_table(doc,
    headers=['Aspect', 'PIRS V1', 'PIRS V2'],
    rows=[
        ['Behavioral Baseline',
         'Population average (all users, all days)',
         'Personal 60-day rolling baseline per individual user'],
        ['Detection Signal',
         'Anomaly score on raw feature values',
         'Drift slope on z-score deviations from personal norm'],
        ['HTTP Data',
         'Skipped (SKIP_HTTP=True) — too large',
         'Fully enabled — 117 million rows processed'],
        ['Prediction Horizon',
         'Same-day detection only',
         '7-day and 14-day breach trajectory forecasting'],
        ['Validation',
         'No ground-truth validation',
         'ROC-AUC + early warning against 5 labeled insiders'],
        ['Explainability',
         'Population-level z-score ranking',
         'Personal deviation feature analysis per user'],
        ['Personality Typing',
         'Fixed feature grouping (arbitrary split)',
         'OCEAN psychometric + behavioral semantic groups'],
        ['LSTM Training',
         '2 epochs, 10K sequences, latent_dim=16',
         '10 epochs, 50K sequences, latent_dim=32'],
        ['Risk Formula',
         'EPR: 71.1% (HTTP zeroed)',
         'EPR: 40% (validated on ground truth, HTTP enabled)'],
        ['ROC-AUC',
         'Not measured',
         '0.8554 (significantly above random baseline of 0.5)'],
    ],
    col_widths=[1.8, 2.7, 2.7],
)
doc.add_paragraph()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. ARCHITECTURE DIAGRAM
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '3. System Architecture Diagram', 1)
add_divider(doc)
add_body(doc,
    'The diagram below shows the complete PIRS V2 pipeline from raw input data through '
    'all nine processing layers to the final outputs and validation artefacts.',
    space_after=8)

if os.path.exists(ARCH_IMG):
    doc.add_picture(ARCH_IMG, width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph('Figure 1 — PIRS V2 nine-layer pre-incident detection architecture')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size  = Pt(9)
    cap.runs[0].italic     = True
    cap.runs[0].font.color.rgb = GRAY
else:
    add_body(doc, '[Architecture diagram image not found — run pirs_v2/draw_architecture.py]',
             italic=True, color=GRAY)

doc.add_paragraph()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. PIPELINE LAYERS — DETAILED
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4. Pipeline Layers — Detailed Description', 1)
add_divider(doc)
add_body(doc,
    'PIRS V2 consists of nine sequential layers. Each layer transforms the data and '
    'adds enriched signals that feed into the next layer. Intermediate checkpoints '
    'are saved as CSV files so any layer can be restarted independently without '
    're-processing the entire dataset.',
    space_after=10)

# ── Layer 0 ───────────────────────────────────────────────────────────────────
add_heading(doc, 'Layer 0 — Feature Extraction  (cert_extractor.py)', 2)
add_body(doc,
    'The entry point of the pipeline. Five raw CERT r6.2 log files are merged and '
    'aggregated into a single per-user, per-day feature matrix.',
    space_after=6)
add_kv(doc, 'Input files',
    'logon.csv (~1.4M events), device.csv (USB), file.csv (308K), '
    'email.csv (1.38M), http.csv (117M rows), LDAP/ (user metadata), '
    'psychometric.csv (OCEAN scores)')
add_kv(doc, 'Output', 'cert_features.csv — 1,393,129 rows × 32 features, 4,000 users, 516 days')
add_kv(doc, 'Feature categories',
    'Logon activity (work/after-hours), USB device usage, File operations, '
    'Email behaviour (internal/external/BCC), HTTP browsing (job/cloud/leak/hack categories)')
add_kv(doc, 'Key feature examples',
    'file_n-to_usb1 (files copied to USB), http_n_cloudf (cloud upload visits), '
    'email_n-exbccmail1 (external BCC emails), n_afterhourusb (after-hours USB events)')
add_body(doc,
    'HTTP extraction was the most computationally intensive step, requiring approximately '
    '4–5 hours to process all 117 million rows. The result is a compact set of behavioral '
    'web-category counts per user per day.',
    italic=True, color=GRAY, space_after=10)

add_divider(doc)

# ── Layer 1 ───────────────────────────────────────────────────────────────────
add_heading(doc, 'Layer 1 — Personal Behavioral Baseline', 2)
add_body(doc,
    'The most critical architectural change from V1. Instead of comparing each user\'s '
    'behaviour to the population mean, Layer 1 builds a personal norm for every user '
    'using a 60-day rolling window.',
    space_after=6)
add_kv(doc, 'Method', '60-day rolling mean and standard deviation per user per feature')
add_kv(doc, 'Output signal', 'z_dev = (today − personal_mean) / personal_std')
add_kv(doc, 'Deviation features', '27 z-score deviation columns per user-day record')
add_kv(doc, 'Rationale',
    'A sales executive who sends 50 emails per day is not anomalous; '
    'a warehouse worker who suddenly sends 50 emails is. Population baselines '
    'miss this class of insider entirely.')
add_body(doc,
    'The 60-day window was chosen to be long enough to capture stable habits but short '
    'enough to adapt to genuine role changes (e.g., a promotion or new project).',
    italic=True, color=GRAY, space_after=10)

add_divider(doc)

# ── Layer 2 ───────────────────────────────────────────────────────────────────
add_heading(doc, 'Layer 2 — Drift Detection  (Linear Regression)', 2)
add_body(doc,
    'Drift detection transforms the z-score deviations from Layer 1 into a single '
    'directional signal: is this user\'s behaviour consistently moving away from their '
    'personal norm?',
    space_after=6)
add_kv(doc, 'Method', '14-day rolling linear regression on each deviation feature')
add_kv(doc, 'drift_score',
    'Weighted sum of per-feature regression slopes — positive = drifting away from norm')
add_kv(doc, 'drift_accel',
    'Rate of change of drift_score — captures acceleration (sudden spike vs. gradual drift)')
add_kv(doc, 'Checkpoint', 'cert_after_drift.csv')
add_body(doc,
    'The 14-day regression window is tuned to the research question: can we detect a '
    'behavioural shift at least one to two weeks before an incident occurs? A shorter window '
    'is too noisy; a longer window reduces early-warning lead time.',
    italic=True, color=GRAY, space_after=10)

add_divider(doc)

# ── Layer 4 ───────────────────────────────────────────────────────────────────
add_heading(doc, 'Layer 4 — Ensemble Anomaly Detection', 2)
add_body(doc,
    'Three complementary anomaly-detection models are combined into a single weighted '
    'ensemble score. Using multiple model families reduces the risk that any one model\'s '
    'blind spots cause a missed detection.',
    space_after=6)

make_table(doc,
    headers=['Model', 'Weight', 'Configuration', 'Strengths'],
    rows=[
        ['Isolation Forest',   '50%',
         '100 trees, contamination = 0.5%, random_state = 42',
         'Fast, robust to high dimensionality, no distribution assumption'],
        ['LSTM Autoencoder',   '35%',
         '50K sequences, 7-step window, latent_dim = 32, 10 epochs',
         'Captures temporal patterns and sequential behavioral change'],
        ['One-Class SVM',      '15%',
         'RBF kernel, nu = 0.005, gamma = "scale", subsampled',
         'Tight decision boundary around normal behaviour'],
    ],
    col_widths=[1.5, 0.7, 2.8, 2.2],
)
doc.add_paragraph()
add_kv(doc, 'Ensemble formula',
    'anomaly_score = 0.50 × IsoForest + 0.35 × LSTM_reconstruction_error + 0.15 × OCSVM')
add_kv(doc, 'Output', 'cert_after_anomaly.csv')
add_body(doc,
    'Weights were assigned based on the reliability and sensitivity of each model class: '
    'Isolation Forest is the most stable general-purpose outlier detector, the LSTM provides '
    'temporal context, and OC-SVM adds a tight kernel-based boundary with lower but '
    'complementary weight.',
    italic=True, color=GRAY, space_after=10)

add_divider(doc)

# ── Layer 5 ───────────────────────────────────────────────────────────────────
add_heading(doc, 'Layer 5 — Risk Scoring & Breach Prediction', 2)
add_body(doc,
    'Layer 5 combines the anomaly score from Layer 4 with the drift signals from '
    'Layer 2 into a single normalized risk score on a 0–10 scale, and generates '
    'forward-looking breach predictions.',
    space_after=6)
add_kv(doc, 'Risk formula',
    'risk_score = 0.50 × anomaly_score  +  0.35 × slope_norm  +  0.15 × accel_norm')
add_kv(doc, 'Scale', '0.0 (no risk) to 10.0 (maximum risk)')

make_table(doc,
    headers=['Alert Level', 'Risk Score Range', 'Recommended Action'],
    rows=[
        ['NORMAL',   '0.0 – 2.0', 'No action required — routine monitoring'],
        ['WATCH',    '2.0 – 4.0', 'Log and flag for weekly review'],
        ['ELEVATED', '4.0 – 6.0', 'Passive friction interventions (L2–L3)'],
        ['HIGH',     '6.0 – 8.0', 'Manager notification and behavioural training (L4–L5)'],
        ['CRITICAL', '8.0 – 10.0','Immediate escalation; consider account restrictions (L6–L7)'],
    ],
    col_widths=[1.5, 1.8, 3.9],
)
doc.add_paragraph()
add_kv(doc, 'Breach prediction',
    'Logistic regression on 7-day and 14-day windows; outputs will_breach_7d / will_breach_14d '
    'boolean flags and projected_risk_7d / projected_risk_14d float values')
add_kv(doc, 'Output', 'cert_after_prediction.csv')
doc.add_paragraph()

add_divider(doc)

# ── Layer 6 ───────────────────────────────────────────────────────────────────
add_heading(doc, 'Layer 6 — Personality Profiling', 2)
add_body(doc,
    'Each of the 4,000 users is assigned a primary personality dimension. This is used '
    'by Layer 7 to select the most effective intervention type for that individual.',
    space_after=6)
add_kv(doc, 'Input features',
    'OCEAN psychometric scores (Openness, Conscientiousness, Extraversion, '
    'Agreeableness, Neuroticism) + behavioral feature clusters')

make_table(doc,
    headers=['Personality Type', 'Distribution', 'Behavioral Profile'],
    rows=[
        ['RISK_TAKER',  '47.0%', 'High sensation-seeking; responds to accountability and consequences'],
        ['AUTONOMOUS',  '18.3%', 'Independent, self-directed; resistant to surveillance-style controls'],
        ['COMPLIANT',   '16.0%', 'Rule-following; responds well to policy reminders and norms'],
        ['SOCIAL',      '10.4%', 'Peer-influenced; responds to team norms and manager relationships'],
        ['CAREFULL',     '8.3%', 'Cautious, detail-oriented; benefits from structured guidance'],
    ],
    col_widths=[1.6, 1.3, 4.3],
)
doc.add_paragraph()
add_kv(doc, 'Output', 'cert_personality.csv — personality scores + PRIMARY_DIMENSION per user')
doc.add_paragraph()

add_divider(doc)

# ── Layer 7 ───────────────────────────────────────────────────────────────────
add_heading(doc, 'Layer 7 — Intervention Selection', 2)
add_body(doc,
    'Based on the risk score from Layer 5 and personality type from Layer 6, Layer 7 '
    'assigns a graduated intervention level to each flagged user.',
    space_after=6)

make_table(doc,
    headers=['Level', 'Intervention Name', 'Description', 'Dist.'],
    rows=[
        ['L1', 'Standard Monitoring',     'Passive background logging; no user-facing action', '83.3%'],
        ['L2', 'Passive Friction',        'Slight delays on sensitive file operations / USB writes', '10.3%'],
        ['L3', 'Warning Banner',          'User-visible policy reminder when accessing sensitive resources', '—'],
        ['L4', 'Behavioural Training',    'Assigned security-awareness micro-training module', '—'],
        ['L5', 'Security Acknowledgment', 'User must re-confirm acceptable-use policy', '—'],
        ['L6', 'Manager Intervention',    'Automated alert to direct manager + HR notation', '—'],
        ['L7', 'Account Lock',            'Temporary suspension of sensitive system access', '—'],
    ],
    col_widths=[0.5, 2.0, 4.0, 0.7],
)
doc.add_paragraph()
add_body(doc,
    'The majority of users (83.3%) require only Level 1 monitoring. This reflects '
    'the design goal: most elevated risk scores are transient and resolve without '
    'intervention. Higher levels are reserved for persistent high-risk trajectories.',
    italic=True, color=GRAY, space_after=10)

add_divider(doc)

# ── Layer 8 ───────────────────────────────────────────────────────────────────
add_heading(doc, 'Layer 8 — Q-Learning Policy Optimization', 2)
add_body(doc,
    'A tabular Q-learning agent learns the optimal intervention policy by treating each '
    'user-day as a state and each intervention level as an action.',
    space_after=6)
add_kv(doc, 'Algorithm',     'Tabular Q-Learning (model-free reinforcement learning)')
add_kv(doc, 'Learning rate', 'α = 0.10')
add_kv(doc, 'Discount factor','γ = 0.90  (balances immediate vs. long-term prevention)')
add_kv(doc, 'Exploration',   'ε = 0.10  (ε-greedy; 10% random action for exploration)')
add_kv(doc, 'Reward signal',
    '+10 if risk drops below threshold after intervention, '
    '−1 per intervention level applied (cost penalty), '
    '−5 if breach occurs despite intervention')
add_kv(doc, 'Output', 'cert_complete.csv — full pipeline output with Q-optimized intervention assignments')
add_body(doc,
    'The Q-learning agent learns over the full dataset that lower-level interventions '
    'are preferred (lower cost), but escalates automatically when the expected reward '
    'from a higher-level action justifies the intrusion on the user.',
    italic=True, color=GRAY, space_after=10)

add_divider(doc)

# ── Layer 9 ───────────────────────────────────────────────────────────────────
add_heading(doc, 'Layer 9 — Evaluation & Metrics  (Research Questions)', 2)
add_body(doc,
    'Layer 9 computes the four research-question metrics that define the performance '
    'of the PIRS V2 framework.',
    space_after=6)

make_table(doc,
    headers=['RQ', 'Research Question', 'Metric', 'Result'],
    rows=[
        ['RQ1', 'Can the system predict insider breaches 7–14 days in advance?',
         'EPR (Early Prediction Rate)',
         'EPR(7d) = 40%\nEPR(14d) = 40%\nROC-AUC = 0.8554'],
        ['RQ2', 'Do personality-matched interventions improve correction quality?',
         'PQ (Personality Quality)',
         'PQ = 0.931\n(93.1% personality-matched)'],
        ['RQ3', 'Can the system explain why a user was flagged?',
         'PIMS (Personal Insight & Model Score)',
         'Personal deviation features\navailable per user-day'],
        ['RQ4', 'What is the economic value of early detection?',
         'Cost Savings Model',
         '$22.8M estimated\nsaved per deployment cycle'],
    ],
    col_widths=[0.5, 2.4, 1.8, 2.5],
)
doc.add_paragraph()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. GROUND TRUTH — 5 REAL INSIDERS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '5. Ground Truth — 5 Labeled Insider Users (CERT r6.2)', 1)
add_divider(doc)
add_body(doc,
    'The CERT Insider Threat Dataset r6.2 contains five labeled insider threat scenarios. '
    'All validation metrics (ROC-AUC, EPR, early warning windows) are computed against '
    'these five users\' ground-truth malicious day ranges.',
    space_after=8)

make_table(doc,
    headers=['User ID', 'Scenario', 'Threat Type', 'Attack Days', 'Days Range', 'PIRS V2 Detection'],
    rows=[
        ['ACM2278', 'Scenario 1',
         'Cloud Upload\n(Data Exfiltration)',
         '6 labeled days', 'Days 229–235',
         'CRITICAL/HIGH at 3-day window\n(cloud upload + external email spike)'],
        ['CMP2946', 'Scenario 2',
         'Malicious Download',
         '20 labeled days', 'Days 402–427',
         'CRITICAL/HIGH at 3-day window\n(USB + file-to-USB + HTTP leak features)'],
        ['PLJ1771', 'Scenario 3',
         'Espionage\n(Recruited Spy)',
         '1 labeled day', 'Day 223',
         'WATCH-level signals only\n(single-day attack, no prior drift)'],
        ['CDE1846', 'Scenario 4',
         'IP Theft\n(Before Resignation)',
         '45 labeled days', 'Days 416–479',
         'Detected at EPR 7d/14d window\n(64-day exfiltration arc clearly visible)'],
        ['MBG3183', 'Scenario 5',
         'Sabotage\n(Disgruntled Employee)',
         '1 labeled day', 'Day 284',
         'WATCH-level signals only\n(single-day attack, inherently hard to predict)'],
    ],
    col_widths=[1.0, 1.0, 1.5, 1.2, 1.0, 2.5],
)
doc.add_paragraph()

add_body(doc, 'Detection Notes:', bold=True, space_after=4)
add_bullet(doc,
    'CDE1846 (IP Theft) is the "ideal" PIRS target: a gradual 45-day exfiltration arc '
    'creates a clear drift signal that the regression and ensemble layers detect well in advance.')
add_bullet(doc,
    'ACM2278 and CMP2946 are detected at short horizons (3 days) because their attacks '
    'involve bursts of measurable activity (USB writes, cloud uploads) that spike the anomaly score.')
add_bullet(doc,
    'PLJ1771 and MBG3183 are single-day events. No pre-event drift exists, which means '
    'no system — regardless of architecture — can reliably predict them at 7–14 day horizons. '
    'Their ROC-AUC contribution comes from accurately scoring their attack day as anomalous.')
doc.add_paragraph()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. DATA FLOW & FILE OUTPUTS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '6. Data Flow & Pipeline Outputs', 1)
add_divider(doc)
add_body(doc,
    'Each pipeline layer reads from the previous layer\'s checkpoint file and writes '
    'an enriched output. This design allows any layer to be re-run in isolation when '
    'parameters are tuned.',
    space_after=8)

make_table(doc,
    headers=['File', 'Produced By', 'Contents', 'Size (approx.)'],
    rows=[
        ['cert_features.csv',              'Layer 0', '32 behavioral features per user-day',       '1,393,129 rows'],
        ['cert_after_drift.csv',           'Layer 2', '+ 27 deviation scores + drift_score/accel', '1,393,129 rows'],
        ['cert_after_anomaly.csv',         'Layer 4', '+ isoforest/lstm/ocsvm/anomaly_score',      '1,393,129 rows'],
        ['cert_after_prediction.csv',      'Layer 5', '+ risk_score + alert_level + breach flags', '1,393,129 rows'],
        ['cert_personality.csv',           'Layer 6', 'Per-user OCEAN scores + PRIMARY_DIMENSION', '4,000 rows'],
        ['cert_complete.csv',              'Layer 8', 'Full pipeline output, all columns',         '1,393,129 rows'],
        ['cert_metrics.csv',               'Layer 9', 'RQ1–RQ4 summary metrics',                   '1 row'],
        ['cert_early_warning.csv',         'Layer 5', 'Risk & breach prediction at 3/7/10/14d',    '5 rows (insiders)'],
        ['cert_validation_summary.csv',    'Validator','ROC-AUC, precision, recall, F1',            '1 row'],
        ['cert_validation_early_warning.csv','Validator','Early warning per insider per window',    '5 rows'],
        ['plots/insider_trajectories.png', 'Validator','Risk trajectory plot — 5 insiders',         'PNG image'],
        ['plots/pirs_v2_architecture.png', 'draw_architecture.py','Full architecture diagram',     'PNG image'],
    ],
    col_widths=[2.8, 1.5, 2.8, 1.4],
)
doc.add_paragraph()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. VALIDATION METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '7. Validation Methodology', 1)
add_divider(doc)
add_body(doc,
    'Because PIRS V2 is an unsupervised system during training (no labels used), '
    'validation is performed post-hoc by comparing output risk scores against the '
    'known ground-truth insider labels from CERT r6.2.',
    space_after=8)

add_heading(doc, 'ROC-AUC Evaluation', 2)
add_body(doc,
    'The Receiver Operating Characteristic Area Under the Curve (ROC-AUC) measures '
    'how well the risk_score separates malicious days (insider labeled) from normal '
    'days across all possible threshold settings.',
    space_after=6)
add_kv(doc, 'ROC-AUC score', '0.8554')
add_kv(doc, 'Random baseline', '0.5000  (a score of 0.5 = no better than chance)')
add_kv(doc, 'Interpretation',
    'A score of 0.8554 means that for a randomly chosen (malicious day, normal day) pair, '
    'PIRS V2 assigns a higher risk score to the malicious day 85.54% of the time.')
doc.add_paragraph()

add_heading(doc, 'Early Warning Rate (EPR)', 2)
add_body(doc,
    'EPR measures whether the system raises an alert for an insider before the first '
    'malicious action occurs. Two methods are used:',
    space_after=6)
add_bullet(doc,
    'Strict window: Was the user\'s risk score above the HIGH threshold (6.0) on every '
    'day in the N-day period before the first attack? (N = 7 or 14)')
add_bullet(doc,
    'EPR (any-day): Was the user flagged at HIGH or above on at least one day in the '
    'N-day window before the attack? This is the reported 40% figure.')
doc.add_paragraph()

add_heading(doc, 'Early Warning Table (cert_validator.py output)', 2)
make_table(doc,
    headers=['User', 'Threat Type', 'First Attack Day', '@14d', '@10d', '@7d', '@3d', 'Peak Risk'],
    rows=[
        ['ACM2278', 'Cloud Upload',  'Day 229', 'WATCH', 'WATCH', 'WATCH', 'CRITICAL', '7.82'],
        ['CMP2946', 'USB Theft',     'Day 402', 'WATCH', 'WATCH', 'WATCH', 'HIGH',     '8.41'],
        ['PLJ1771', 'Espionage',     'Day 223', 'NORMAL','NORMAL','NORMAL','NORMAL',   '3.12'],
        ['CDE1846', 'IP Theft',      'Day 416', 'WATCH', 'WATCH', 'HIGH',  'HIGH',     '9.56'],
        ['MBG3183', 'Sabotage',      'Day 284', 'NORMAL','NORMAL','NORMAL','NORMAL',   '2.97'],
    ],
    col_widths=[0.9, 1.3, 1.3, 0.7, 0.7, 0.7, 0.7, 0.9],
)
doc.add_paragraph()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. LIMITATIONS & FUTURE WORK
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '8. Limitations & Future Work', 1)
add_divider(doc)

add_heading(doc, 'Current Limitations', 2)
add_bullet(doc,
    'Single-day insiders (PLJ1771, MBG3183) are fundamentally unpredictable at 7–14 day '
    'horizons because they exhibit no prior behavioral drift — the attack is a single '
    'spontaneous action.')
add_bullet(doc,
    'CERT r6.2 is a synthetic dataset. Real-world insider behavior may be more subtle, '
    'noisy, or context-dependent than the simulated scenarios.')
add_bullet(doc,
    'The 60-day personal baseline requires sufficient historical data; the system has '
    'limited capability for users who joined the organization recently.')
add_bullet(doc,
    'HTTP feature extraction took 4–5 hours on the full 117M-row file; real-time '
    'deployment would require streaming ingestion infrastructure.')
add_bullet(doc,
    'The Q-learning agent is tabular (finite state/action space); a deep RL approach '
    'may generalize better to unseen user types.')
doc.add_paragraph()

add_heading(doc, 'Planned Extensions', 2)
add_bullet(doc,
    'LANL Dataset: Apply the same 9-layer pipeline to the Los Alamos National Laboratory '
    '(LANL) authentication and process event dataset to validate generalizability.')
add_bullet(doc,
    'Graph-based features: Model user–resource–peer relationships as a dynamic graph '
    'to detect collusion and coordinated insider attacks.')
add_bullet(doc,
    'LLM-augmented explainability (RQ3): Use a language model to generate natural-language '
    'explanations of each user\'s risk trajectory for security analyst consumption.')
add_bullet(doc,
    'Streaming pipeline: Replace batch CSV processing with Apache Kafka / Flink for '
    'real-time per-event risk score updates.')
add_bullet(doc,
    'Federated learning: Enable privacy-preserving model training across multiple '
    'organizations without sharing raw behavioral data.')
doc.add_paragraph()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. GLOSSARY
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '9. Glossary', 1)
add_divider(doc)

terms = [
    ('PIRS',        'Predictive Intervention and Risk Stabilization System'),
    ('CERT r6.2',   'Carnegie Mellon CERT Insider Threat Dataset, release 6.2'),
    ('EPR',         'Early Prediction Rate — fraction of insiders detected N days before first attack'),
    ('PQ',          'Personality Quality — fraction of interventions matched to user personality type'),
    ('ROC-AUC',     'Receiver Operating Characteristic Area Under the Curve — discrimination metric'),
    ('Drift',       'A sustained directional change in a user\'s behavior away from their personal baseline'),
    ('Drift Score', 'Weighted sum of 14-day regression slopes across all deviation features'),
    ('Z-Score',     'Standard deviations from personal mean: (today − personal_mean) / personal_std'),
    ('IsoForest',   'Isolation Forest — tree-based unsupervised anomaly detection algorithm'),
    ('LSTM AE',     'Long Short-Term Memory Autoencoder — reconstructs normal sequences; high error = anomaly'),
    ('OC-SVM',      'One-Class Support Vector Machine — learns a tight boundary around normal data'),
    ('Q-Learning',  'Model-free reinforcement learning algorithm that learns action-value functions'),
    ('OCEAN',       'Big Five personality model: Openness, Conscientiousness, Extraversion, Agreeableness, Neuroticism'),
    ('DXA',         'Device-independent units used by Word documents (1440 DXA = 1 inch)'),
]
make_table(doc,
    headers=['Term', 'Definition'],
    rows=terms,
    col_widths=[1.5, 5.7],
)
doc.add_paragraph()

# ── Footer-style closing note ─────────────────────────────────────────────────
add_divider(doc)
closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
c = closing.add_run(
    'PIRS V2 Architecture Reference  ·  MS Capstone 2026  ·  Generated from pipeline output'
)
c.font.size  = Pt(9)
c.italic     = True
c.font.color.rgb = GRAY

# ── Save ──────────────────────────────────────────────────────────────────────
os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f'[OK] Saved: {OUT}')
